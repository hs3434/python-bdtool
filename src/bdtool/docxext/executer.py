from __future__ import annotations
from typing import TYPE_CHECKING
from docx.shared import RGBColor
from bdtool.tool import common_color_rgb
from bdtool.docxext.math import add_math, parse_math
from bdtool.docxext.endnote import add_endnote
if TYPE_CHECKING:
    from docx.styles.style import ParagraphStyle, CharacterStyle
    from docx.document import Document
    from docx.shared import Length


class ClassTool:
    @classmethod
    def try_call(self, obj):
        if callable(obj):
            return obj()
        else:
            return obj

    @classmethod
    def create_call(self, func, *args, **kwarg):
        def fun():
            return func(*args, **kwarg)
        return fun


class Executer(ClassTool):
    def __init__(self):
        self.tasks = []
        self.nodes = []
        self.sub_parts: list[Executer] = []

    def execute(self, doc: Document):
        for task in self.get_tasks():
            if task["type"] == "paragraph":
                style = task.get("style", None)
                text = task.get("text", "")
                obj = doc.add_paragraph(text, style=style)
            elif task["type"] == "run":
                style = task.get("style", None)
                text = task.get("text", "")
                run = obj.add_run(text, style=style)
                if task.get("bold", None) is not None:
                    run.bold = task["bold"]
                if task.get("italic", None) is not None:
                    run.italic = task["italic"]
                if task.get("color", None) is not None:
                    run.font.color.rgb = RGBColor(*common_color_rgb(task["color"]))
            elif task["type"] == "picture":
                path = task.get("path", "")
                width = task.get("width", None)
                height = task.get("height", None)
                style = task.get("style", obj.style)
                doc.add_picture(path, width=width, height=height)
                obj = doc.add_paragraph(style=style)
            elif task["type"] == "math":
                style = task.get("style", None)
                text = task.get("text", "")
                obj = doc.add_paragraph(style=style)
                add_math(obj._element, text)
            elif task["type"] == "endnote":
                text = task.get("text", "")
                index = task.get("index", "1")
                add_endnote(obj._element, text, index)
            else:
                raise NotImplementedError("no type named {}".format(task["type"]))

    def get_tasks(self):
        tasks = self.tasks.copy()
        for sub_part in self.sub_parts:
            tasks.extend(sub_part.get_tasks())
        return tasks
    
    def sub_part(self):
        part = Executer()
        self.sub_parts.append(part)
        return part

    def add_paragraph(self, text: str = "", style: str | ParagraphStyle | None = None, 
                      color: str | None = None, bold: bool | None = None, italic: bool | None = None):
        self.tasks.append({
            "type": "paragraph",
            "text": text,
            "style": style,
            "color": color,
            "bold": bold,
            "italic": italic
        })
    
    def add_run(self, text: str | None = None, style: str | CharacterStyle | None = None,
                color: str | None = None, bold: bool | None = None, italic: bool | None = None):
        self.tasks.append({
            "type": "run",
            "text": text,
            "style": style,
            "color": color,
            "bold": bold,
            "italic": italic
        })

    def add_math(self, latex: str, style: str | ParagraphStyle | None = None):
        self.tasks.append({
            "type": "math",
            "text": latex,
            "style": style
        })
        
    def add_endnote(self, instrText: str, index: int | None = None):
        self.tasks.append({
            "type": "endnote",
            "text": instrText,
            "index": index
        })
        
    def add_picture(self, path: str,
                    width: int | Length | None = None,
                    height: int | Length | None = None,
                    style: str | ParagraphStyle | None = None):
        self.tasks.append({
            "type": "picture",
            "path": path,
            "width": width,
            "height": height,
            "style": style
        })